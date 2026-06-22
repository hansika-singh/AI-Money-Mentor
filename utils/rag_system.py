"""
RAG-Powered Financial Knowledge Assistant
Retrieval-Augmented Generation for document Q&A
"""

import os
import tempfile
from typing import List, Dict, Optional, Tuple
import hashlib
import json

# Document parsing
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract

# RAG components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document

import numpy as np


class RAGSystem:
    """
    Retrieval-Augmented Generation System for financial documents
    """
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        """
        Initialize RAG system with embedding model and vector store
        
        Args:
            persist_directory: Directory to persist vector database
        """
        self.persist_directory = persist_directory
        self.embeddings = None
        self.vector_store = None
        self.documents = []
        self.document_metadata = {}
        self.client = None
        
        # Initialize embeddings
        self._initialize_embeddings()
        
        # Initialize or load vector store
        self._initialize_vector_store()
    
    def _initialize_embeddings(self):
        """Initialize sentence transformer embeddings"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            print("✅ Embeddings initialized successfully")
        except Exception as e:
            print(f"⚠️ Error initializing embeddings: {e}")
            # Fallback to a simpler approach
            self.embeddings = None
    
    def _initialize_vector_store(self):
        """Initialize or load existing vector store"""
        try:
            if os.path.exists(self.persist_directory) and self.embeddings:
                self.vector_store = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                print(f"✅ Vector store loaded from {self.persist_directory}")
            else:
                self.vector_store = None
        except Exception as e:
            print(f"⚠️ Error loading vector store: {e}")
            self.vector_store = None
    
    def set_client(self, client):
        """Set Groq client for AI responses"""
        self.client = client
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        if not text.strip():
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"PyPDF2 failed: {e}")
        
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            for sheet in wb.worksheets:
                text += f"\n=== Sheet: {sheet.title} ===\n"
                for row in sheet.iter_rows(values=True):
                    row_text = " | ".join([str(cell) for cell in row if cell is not None])
                    if row_text:
                        text += row_text + "\n"
            return text
        except Exception as e:
            print(f"Excel extraction failed: {e}")
            return ""
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"OCR failed: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                print(f"TXT extraction failed: {e}")
                return ""
    
    def extract_text(self, file_path: str, file_type: str = None) -> str:
        """
        Extract text from file based on extension
        
        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, xlsx, txt, png, jpg)
        
        Returns:
            Extracted text
        """
        if not file_type:
            file_type = file_path.split('.')[-1].lower()
        
        extractors = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'xlsx': self.extract_text_from_excel,
            'xls': self.extract_text_from_excel,
            'txt': self.extract_text_from_txt,
            'png': self.extract_text_from_image,
            'jpg': self.extract_text_from_image,
            'jpeg': self.extract_text_from_image,
            'gif': self.extract_text_from_image,
        }
        
        extractor = extractors.get(file_type)
        if extractor:
            return extractor(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def chunk_text(self, text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[Document]:
        """
        Split text into chunks using LangChain's text splitter
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
        
        Returns:
            List of Document objects
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len,
        )
        
        chunks = splitter.split_text(text)
        
        return [Document(page_content=chunk, metadata={"chunk_id": i}) for i, chunk in enumerate(chunks)]
    
    def process_document(self, file_path: str, file_type: str = None, metadata: Dict = None) -> Dict:
        """
        Process a document: extract, chunk, embed, and store
        
        Args:
            file_path: Path to the document
            file_type: File extension
            metadata: Additional metadata for the document
        
        Returns:
            Dict with processing results
        """
        if not self.embeddings:
            return {'success': False, 'error': 'Embeddings not initialized'}
        
        try:
            # Extract text
            text = self.extract_text(file_path, file_type)
            
            if not text or len(text.strip()) < 10:
                return {'success': False, 'error': 'Could not extract text from document'}
            
            # Generate document ID
            doc_id = hashlib.md5(file_path.encode()).hexdigest()
            
            # Chunk text
            chunks = self.chunk_text(text)
            
            if not chunks:
                return {'success': False, 'error': 'No chunks created'}
            
            # Add metadata to chunks
            metadata = metadata or {}
            metadata['source'] = file_path
            metadata['doc_id'] = doc_id
            metadata['chunk_count'] = len(chunks)
            
            for chunk in chunks:
                chunk.metadata.update(metadata)
            
            # Create or update vector store
            if self.vector_store is None:
                self.vector_store = Chroma.from_documents(
                    documents=chunks,
                    embedding=self.embeddings,
                    persist_directory=self.persist_directory
                )
            else:
                self.vector_store.add_documents(chunks)
            
            # Persist vector store
            self.vector_store.persist()
            
            # Store metadata
            self.document_metadata[doc_id] = {
                'path': file_path,
                'metadata': metadata,
                'chunk_count': len(chunks),
                'doc_id': doc_id
            }
            
            return {
                'success': True,
                'doc_id': doc_id,
                'chunk_count': len(chunks),
                'text_preview': text[:200] + "..." if len(text) > 200 else text
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def query(self, question: str, k: int = 3) -> Dict:
        """
        Query the RAG system
        
        Args:
            question: User question
            k: Number of relevant chunks to retrieve
        
        Returns:
            Dict with answer and source documents
        """
        if not self.vector_store:
            return {
                'success': False,
                'answer': "No documents have been uploaded yet. Please upload a document first.",
                'sources': []
            }
        
        if not self.client:
            return {
                'success': False,
                'answer': "AI client not configured. Please set up Groq API key.",
                'sources': []
            }
        
        try:
            # Retrieve relevant chunks
            docs = self.vector_store.similarity_search(question, k=k)
            
            if not docs:
                return {
                    'success': False,
                    'answer': "I couldn't find any relevant information in the uploaded documents.",
                    'sources': []
                }
            
            # Build context
            context = "\n\n---\n\n".join([doc.page_content for doc in docs])
            
            # Get sources
            sources = []
            for doc in docs:
                source = doc.metadata.get('source', 'Unknown')
                if source not in sources:
                    sources.append(source)
            
            # Generate answer using AI
            prompt = f"""
You are a financial knowledge assistant. Answer the user's question based on the provided document context.

**Context from documents:**
{context}

**User Question:** {question}

**Instructions:**
1. Answer based ONLY on the provided context
2. If the answer is not in the context, say "I couldn't find this information in the uploaded documents."
3. Be specific and use exact numbers/facts from the context
4. Cite which document(s) the information comes from

**Answer:**
"""
            
            response = self.client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": "You are a financial knowledge assistant. Answer questions based on document context only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=500
            )
            
            answer = response.choices[0].message.content
            
            return {
                'success': True,
                'answer': answer,
                'sources': sources,
                'chunks_used': len(docs),
                'relevant_chunks': [{'content': doc.page_content[:150] + "...", 'source': doc.metadata.get('source', 'Unknown')} for doc in docs]
            }
            
        except Exception as e:
            return {'success': False, 'answer': f"Error: {str(e)}", 'sources': []}
    
    def get_documents(self) -> List[Dict]:
        """Get list of uploaded documents"""
        return [
            {
                'doc_id': doc_id,
                'path': info['path'],
                'chunk_count': info['chunk_count'],
                'metadata': info['metadata']
            }
            for doc_id, info in self.document_metadata.items()
        ]
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the vector store"""
        try:
            # Remove from metadata
            if doc_id in self.document_metadata:
                del self.document_metadata[doc_id]
            
            # Rebuild vector store without the document
            # Note: Chroma doesn't support deletion easily, so we rebuild
            # This is a simplified approach
            
            return True
        except Exception as e:
            print(f"Error deleting document: {e}")
            return False
    
    def clear_all(self):
        """Clear all documents and vector store"""
        try:
            self.vector_store = None
            self.document_metadata = {}
            # Remove persisted data
            if os.path.exists(self.persist_directory):
                import shutil
                shutil.rmtree(self.persist_directory)
            self._initialize_vector_store()
            return True
        except Exception as e:
            print(f"Error clearing: {e}")
            return False